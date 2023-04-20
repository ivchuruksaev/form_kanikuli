from flask import Flask, render_template, request, redirect, url_for, send_file
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
import os
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user
from docx2pdf import convert
import pythoncom
from docx import Document
import comtypes.client
import time
from flask import after_this_request
import pymorphy2

morph = pymorphy2.MorphAnalyzer()

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "your-secret-key")

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://tester:pass@localhost:3307/mydatabase'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

def convert_docx_to_pdf(temp_docx, output_file):
    word_app = comtypes.client.CreateObject('Word.Application')
    word_app.Visible = False
    doc = word_app.Documents.Open(temp_docx)
    doc.SaveAs(output_file, FileFormat=17)  # 17 is the code for PDF format
    doc.Close()
    word_app.Quit()

@app.route('/update_parent_info', methods=['POST'])
def update_parent_info():
    parent_id = request.form.get('parent_id')
    parent = ParentData.query.get(parent_id)

    if parent:
        parent.parent_name_rp = request.form['parent_name_rp']
        parent.child_name_ip = request.form['child_name_ip']
        parent.child_name_rp = request.form['child_name_rp']
        parent.child_name_vp = request.form['child_name_vp']
        parent.child_name_dp = request.form['child_name_dp']
        parent.child_name_tp = request.form['child_name_tp']
        parent.contract_start = request.form['contract_start']
        parent.contract_end = request.form['contract_end']

        db.session.commit()

    return redirect(url_for('admin'))

def inflect_pymorphy(text, case, gender=None):
    words = text.split()
    inflected_words = []

    for word in words:
        parsed_word = morph.parse(word)[0]
        if gender:
            inflected_word = parsed_word.inflect({case, gender})
        else:
            inflected_word = parsed_word.inflect({case})
        if inflected_word is not None:
            inflected_words.append(inflected_word.word.capitalize())
        else:
            inflected_words.append(word.capitalize())

    return " ".join(inflected_words)

import re
def replace_text_in_docx(docx_path, replacements):
    doc = Document(docx_path)

    # Replace text in paragraphs
    for p in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if old_text in inline[i].text:
                        text = inline[i].text.replace(old_text, new_text)
                        inline[i].text = text

    # Replace text in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in p.text:
                            print(f"Cell Paragraph - Before: {p.text}")
                            text = "".join(run.text for run in p.runs)
                            text = re.sub(old_text, new_text, text)
                            i = 0
                            for run in p.runs:
                                run_length = len(run.text)
                                run.text = text[i:i + run_length]
                                i += run_length
                            print(f"Cell Paragraph - After: {p.text}")

    return doc

class User(UserMixin):
    def __init__(self, id):
        self.id = id

class ParentData(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    parent_last_name = db.Column(db.String(100), nullable=False)
    parent_first_name = db.Column(db.String(100), nullable=False)
    parent_patronymic = db.Column(db.String(100))
    child_last_name = db.Column(db.String(100), nullable=False)
    child_first_name = db.Column(db.String(100), nullable=False)
    child_patronymic = db.Column(db.String(100))
    child_birth_date = db.Column(db.Date, nullable=False)
    child_gender = db.Column(db.String(10), nullable=False)
    phone = db.Column(db.String(20), nullable=False)
    residence_info = db.Column(db.Text, nullable=False)
    medical_agreement = db.Column(db.Boolean, nullable=False)
    leisure_activities_agreement = db.Column(db.Boolean, nullable=False)
    mobile_agreement = db.Column(db.Boolean, nullable=False)
    rules_agreement = db.Column(db.Boolean, nullable=False)
    notOANO_agreement = db.Column(db.Boolean, nullable=False)
    transportation = db.Column(db.Integer, nullable=False)
    timestamp = db.Column(db.DateTime, index=True, default=datetime.utcnow)
    emergency_phone = db.Column(db.String(20), nullable=False)
    registration = db.Column(db.String(100), nullable=False)
    contract_start = db.Column(db.Date, nullable=True)
    contract_end = db.Column(db.Date, nullable=True)
    passport_series = db.Column(db.String(4), nullable=False)
    passport_number = db.Column(db.String(6), nullable=False)
    passport_issued_by = db.Column(db.String(255), nullable=False)
    parent_name_rp = db.Column(db.String(100), nullable=False)
    child_name_rp = db.Column(db.String(100), nullable=False)
    child_name_tp = db.Column(db.String(100), nullable=False)
    child_name_vp = db.Column(db.String(100), nullable=False)
    child_name_dp = db.Column(db.String(100), nullable=False)
    child_name_ip = db.Column(db.String(100), nullable=False)

@login_manager.user_loader
def load_user(user_id):
    return User(user_id)

@app.route('/admin', methods=['GET', 'POST'])
@login_required
def admin():
    if request.method == 'POST':
        parent_id = request.form['parent_id']
        contract_start = request.form['contract_start']
        contract_end = request.form['contract_end']
        parent = ParentData.query.get(parent_id)
        parent.contract_start = contract_start
        parent.contract_end = contract_end
        db.session.commit()

    parents = ParentData.query.all()
    return render_template('admin.html', parents=parents)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if username == 'admin' and password == 'password':
            user = User(1)
            login_user(user)
            return redirect(url_for('admin'))
        else:
            flash('Неверное имя пользователя или пароль')
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))

def guess_parent_gender(patronymic):
    if patronymic[-1].lower() == 'ч':
        return 'masc'
    elif patronymic[-1].lower() == 'а':
        return 'femn'
    else:
        return 'femn'

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':


        parent_name = f"{request.form['parent_last_name']} {request.form['parent_first_name']} {request.form['parent_patronymic']}"
        parent_gender = guess_parent_gender(request.form['parent_patronymic'])
        parent_name_rp = inflect_pymorphy(parent_name, 'gent', parent_gender)
        child_name = f"{request.form['child_last_name']} {request.form['child_first_name']} {request.form['child_patronymic']}"
        child_gender = request.form['child_gender']
        child_name_rp = inflect_pymorphy(child_name, 'gent', child_gender)
        child_name_tp = inflect_pymorphy(child_name, 'ablt', child_gender)
        child_name_vp = inflect_pymorphy(child_name, 'accs', child_gender)
        child_name_dp = inflect_pymorphy(child_name, 'datv', child_gender)
        child_name_ip = inflect_pymorphy(child_name, 'nomn', child_gender)

        # Обработка данных формы
        parent_data = ParentData(
            parent_name_rp=parent_name_rp,
            child_name_rp=child_name_rp,
            child_name_tp=child_name_tp,
            child_name_vp=child_name_vp,
            child_name_dp=child_name_dp,
            child_name_ip=child_name_ip,
            parent_last_name=request.form['parent_last_name'],
            parent_first_name=request.form['parent_first_name'],
            parent_patronymic=request.form['parent_patronymic'] if not request.form.get('parent_no_patronymic') else None,
            child_last_name=request.form['child_last_name'],
            child_first_name=request.form['child_first_name'],
            child_patronymic=request.form['child_patronymic'] if not request.form.get('child_no_patronymic') else None,
            child_birth_date=request.form['child_birth_date'],
            child_gender=request.form['child_gender'],
            phone=request.form['phone'],
            passport_series=request.form['passport_series'],
            passport_number=request.form['passport_number'],
            residence_info=request.form['residence_info'],
            medical_agreement=request.form.get('medical_agreement') == 'true',
            leisure_activities_agreement=request.form.get('leisure_activities_agreement') == 'true',
            mobile_agreement=request.form.get('mobile_agreement') == 'true',
            rules_agreement=request.form.get('rules_agreement') == 'true',
            notOANO_agreement=request.form.get('notOANO_agreement') == 'true',
            transportation=int(request.form['transportation']),
            emergency_phone=request.form['emergency_phone'],
            passport_issued_by=request.form['passport_issued_by'],
            registration=request.form['registration']

        )

        db.session.add(parent_data)
        db.session.commit()

        return redirect(url_for('success'))
    return render_template('index.html', success=False)

@app.route('/success')
def success():
    return render_template('index.html', success=True)

@app.route('/download_contract/<int:parent_id>')
def download_contract(parent_id):
    parent = ParentData.query.get(parent_id)

    input_docx = 'C:/Users/Ivan/source/repos/Form/Form/templates/template.docx'
    timestamp = str(time.time())
    temp_docx = f'C:/Users/Ivan/source/repos/Form/Form/templates/temp_{timestamp}.docx'
    output_pdf = f'C:/Users/Ivan/source/repos/Form/Form/templates/generated_contract_{timestamp}.pdf'
    contract_start_day = parent.contract_start.strftime("%d")
    contract_start_month = parent.contract_start.strftime("%B")
    contract_end_day = parent.contract_end.strftime("%d")
    contract_end_month = parent.contract_end.strftime("%B")
    # Словарь для перевода английских названий месяцев на русский
    month_translation = {
    "January": "января",
    "February": "февраля",
    "March": "марта",
    "April": "апреля",
    "May": "мая",
    "June": "июня",
    "July": "июля",
    "August": "августа",
    "September": "сентября",
    "October": "октября",
    "November": "ноября",
    "December": "декабря"
    }
    contract_start_month_rus = month_translation[contract_start_month]
    contract_end_month_rus = month_translation[contract_end_month]

    replacements = {
        '11111': parent.parent_last_name,
        '11112': parent.parent_first_name,
        '11113': parent.parent_patronymic,
        '11121': parent.parent_name_rp.split()[0],
        '11122': parent.parent_name_rp.split()[1],
        '11123': parent.parent_name_rp.split()[2],
        '22221': parent.child_name_rp.split()[0],
        '22222': parent.child_name_rp.split()[1],
        '22223': parent.child_name_rp.split()[2],
        '22231': parent.child_name_tp.split()[0],
        '22232': parent.child_name_tp.split()[1],
        '22233': parent.child_name_tp.split()[2],
        '22241': parent.child_name_vp.split()[0],
        '22242': parent.child_name_vp.split()[1],
        '22243': parent.child_name_vp.split()[2],
        '22271': parent.child_name_ip.split()[0],
        '22272': parent.child_name_ip.split()[1],
        '22273': parent.child_name_ip.split()[2],
        '22211': parent.child_name_dp.split()[0],
        '22212': parent.child_name_dp.split()[1],
        '22213': parent.child_name_dp.split()[2],
        '-1': contract_start_day,
        '0987': contract_start_month_rus,
        '-2': contract_end_day,
        '7890': contract_end_month_rus,
        'СУММАЦИФРЫ': '13 500',
        'СУММАБУКВЫ': 'тринадцать тысяч пятьсот',
        'ПИТАНИЕСУММА': '1000',
        '0000': parent.passport_series,
        '777777': parent.passport_number,
        'КОГДАИКЕМВЫДАН': parent.passport_issued_by,
        'МЕСТОЖИТЕЛЬСТВА': parent.residence_info,
        'РЕГИСТРАЦИЯ': parent.registration,
        'НОМЕРТЕЛ': parent.phone,
        'ЭКСТРЕННЫЙТЕЛ': parent.emergency_phone,
        'ДАТАРОЖД': parent.child_birth_date.strftime("%d-%m-%Y")
    }
    print(parent.parent_name_rp.split()[0])
    if parent.medical_agreement:
        replacements['СОГ1'] = '+'
        replacements['СОГ2'] = ' '
    else:
        replacements['СОГ1'] = ''
        replacements['СОГ2'] = '+'
    if parent.leisure_activities_agreement:
        replacements['СОГ3'] = '+'
        replacements['СОГ4'] = ' '
    else:
        replacements['СОГ3'] = ''
        replacements['СОГ4'] = '+'
    if parent.mobile_agreement:
        replacements['СОГ5'] = '+'
        replacements['СОГ6'] = ' '
    else:
        replacements['СОГ5'] = ''
        replacements['СОГ6'] = '+'
    if parent.rules_agreement:
        replacements['СОГ7'] = '+'
        replacements['СОГ8'] = ' '
    else:
        replacements['СОГ7'] = ''
        replacements['СОГ8'] = '+'
    if parent.notOANO_agreement:
        replacements['СОГ9'] = '+'
        replacements['СОГ0'] = ' '
    else:
        replacements['СОГ9'] = ''
        replacements['СОГ0'] = '+'
    if parent.transportation == 1:
        replacements['СОГГ1'] = '+'
        replacements['СОГГ2'] = ' '
        replacements['СОГГ3'] = ' '
    else: 
        if parent.transportation == 2:
            replacements['СОГГ1'] = ' '
            replacements['СОГГ2'] = '+'
            replacements['СОГГ3'] = ' '
        else:
            if parent.transportation == 3:
                replacements['СОГГ1'] = ' '
                replacements['СОГГ2'] = ' '
                replacements['СОГГ3'] = '+'

    
    doc = replace_text_in_docx(input_docx, replacements)
    doc.save(temp_docx)

    pythoncom.CoInitialize()
    convert_docx_to_pdf(temp_docx, output_pdf)
    filename = f"contract_{parent.parent_last_name}_{parent.parent_first_name}_{parent.parent_patronymic}.pdf"

    @after_this_request
    def cleanup(response):
        try:
            os.remove(temp_docx)
            os.remove(output_pdf)
        except Exception as e:
            print("Error deleting temporary files:", e)
        return response


    return send_file(output_pdf, as_attachment=True, download_name=filename)

if __name__ == '__main__':
    app.run(debug=True)